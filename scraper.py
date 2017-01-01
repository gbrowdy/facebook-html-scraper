from bs4 import BeautifulSoup
import json
from docx import Document
from docx.shared import Inches

soup = BeautifulSoup(open("occupation.html"), 'html.parser')


def getTopLevelComments(tag):
    return tag.has_attr('class') and 'UFIComment' in tag['class'] and tag.has_attr('aria-label') and tag[
                                                                                                         'aria-label'] == 'Comment'


def getTopLevelReplies(tag):
    return tag.has_attr('class') and 'UFIComment' in tag['class'] and tag.has_attr('aria-label') and tag[
                                                                                                         'aria-label'] == 'Comment reply'


def getCommentBodyTag(tag):
    return tag.has_attr('class') and 'UFICommentBody' in tag['class']


def getCommentAuthorTag(tag):
    return tag.has_attr('class') and 'UFICommentActorName' in tag.get('class')


def hasReplies(comment):
    return comment.next_sibling.next_sibling and comment.next_sibling.next_sibling.has_attr(
        'class') and 'UFIReplyList' in comment.next_sibling.next_sibling.get('class')


def getCommentBodyText(comment):
    full_text = ""
    body = comment.find(getCommentBodyTag)
    for string in body.strings:
        full_text += unicode(string).splitlines()[0]
    return full_text


def getCommentAuthor(comment):
    tag = comment.find(getCommentAuthorTag)
    return unicode(tag.text).splitlines()[0]


def getReplies(comment):
    head = comment.next_sibling.next_sibling
    return head.find_all(getTopLevelReplies)


class Comment(object):
    def __init__(self, comment):
        self.body = getCommentBodyText(comment)
        self.author = getCommentAuthor(comment)


def serializer(obj):
    return obj.__dict__


top_level_comments = soup.find_all(getTopLevelComments)

parsed_comments = []

for comment in top_level_comments:
    parsed_comments.append({'comment': Comment(comment)})
    if hasReplies(comment):
        replies = getReplies(comment)
        parsed_replies = []
        for reply in replies:
            parsed_replies.append({'reply': Comment(reply)})
        parsed_comments[len(parsed_comments) - 1]['replies'] = parsed_replies

jsonList = json.dumps(parsed_comments, default=serializer)
doc = Document()
for comment in json.loads(jsonList):
    doc.add_heading(comment["comment"]["author"], level=1)
    doc.add_paragraph(comment["comment"]["body"])
    if comment.get("replies"):
        for reply in comment["replies"]:
            p = doc.add_paragraph(reply["reply"]["author"], style="Heading 1")
            p.paragraph_format.left_indent = Inches(0.5)
            p = doc.add_paragraph(reply["reply"]["body"])
            p.paragraph_format.left_indent = Inches(0.5)

doc.save('test.docx')
